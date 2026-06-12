import os
import re
import uuid
import zipfile
import subprocess
import logging
import json
import hmac
import shutil
import time
from datetime import timedelta
from urllib.parse import urlencode
from urllib.request import urlopen
from flask import Flask, request, jsonify, send_file, render_template, redirect, session, url_for
import openpyxl
try:
    import xlrd
    HAS_XLRD = True
except ImportError:
    HAS_XLRD = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.environ.get('APP_SECRET_KEY', 'contract-fill-system-session-key')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=8)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'output'
app.config['DATA_FOLDER'] = os.environ.get('CONTRACT_DATA_FOLDER', '/opt/contract-platform-data')
app.config['TEMPLATE_FOLDER'] = os.environ.get('CONTRACT_TEMPLATE_FOLDER', '/opt/contract-platform-templates/电子合同')
app.config['AMAP_API_KEY'] = os.environ.get('AMAP_API_KEY') or os.environ.get('GAODE_API_KEY') or ''

LOGIN_USERNAME = 'admin'
LOGIN_PASSWORD = 'yhd,123'
LOGIN_ATTEMPTS = {}
LOGIN_MAX_ATTEMPTS = 5
LOGIN_LOCK_SECONDS = 300
IDLE_TIMEOUT_SECONDS = 300
ALLOWED_EXTENSIONS = {'doc', 'docx', 'xls', 'xlsx'}
PRINT_PAGE_SIZE = (1654, 2339)
PRINT_CONTENT_BOX = (85, 145, 1569, 2030)

FIELD_DEFS = [
    {'key': 'plate',         'label': '车牌号（苏K-后面部分）', 'example': '12345'},
    {'key': 'seat_count',    'label': '座位数',             'example': '45'},
    {'key': 'car_type',      'label': '车型',               'example': '大巴'},
    {'key': 'driver1_name',  'label': '驾驶员1姓名',         'example': '张三'},
    {'key': 'driver1_phone', 'label': '驾驶员1电话',         'example': '13800000001'},
    {'key': 'driver1_cert',  'label': '驾驶员1从业资格证号',  'example': 'JS0001'},
    {'key': 'transport_cert','label': '道路运输证号',        'example': 'YZ001'},
    {'key': 'start_date',    'label': '用车开始日期',        'example': '2026-06-15', 'type': 'date'},
    {'key': 'end_date',      'label': '用车结束日期',        'example': '2026-06-16', 'type': 'date'},
    {'key': 'dest',          'label': '目的地',             'example': '南京'},
    {'key': 'route',         'label': '途经',               'example': '宁沪高速'},
    {'key': 'fee',           'label': '运输费用（元）',      'example': '2000'},
    {'key': 'sign_date',     'label': '签订日期',           'example': '2026-06-10', 'type': 'date'},
    {'key': 'driver2_name',  'label': '驾驶员2姓名',         'example': ''},
    {'key': 'driver2_phone', 'label': '驾驶员2电话',         'example': ''},
    {'key': 'driver2_cert',  'label': '驾驶员2从业资格证号',  'example': ''},
]

NANJING_BLANK_MAP = [
    (4,  0, 'seat_count'),
    (4,  1, 'car_type'),
    (4,  3, 'plate'),
    (5,  0, 'driver1_name'),
    (5,  1, 'driver1_phone'),
    (6,  0, 'driver1_cert'),
    (7,  0, 'driver2_name'),
    (7,  1, 'driver2_phone'),
    (8,  0, 'driver2_cert'),
    (9,  1, 'transport_cert'),
    (9,  2, 'start_year'),
    (9,  3, 'start_month'),
    (9,  4, 'start_day'),
    (9,  5, 'end_year'),
    (9,  6, 'end_month'),
    (9,  7, 'end_day'),
    (9,  9, 'dest'),
    (10, 0, 'route'),
    (11, 2, 'fee'),
    (24, 0, 'sign_year'),
    (24, 1, 'sign_month'),
    (24, 2, 'sign_day'),
]

VEHICLE_DATA_FILE = 'vehicles.xlsx'
DRIVER_DATA_FILE = 'drivers.xlsx'
DATA_CACHE = {
    'vehicle_mtime': None,
    'vehicle_records': [],
    'driver_mtime': None,
    'driver_records': [],
}


def read_excel(path):
    ext = path.rsplit('.', 1)[-1].lower()
    if ext == 'xls':
        if not HAS_XLRD:
            raise RuntimeError('服务器缺少xlrd库，无法读取.xls文件')
        wb = xlrd.open_workbook(path)
        ws = wb.sheet_by_index(0)
        all_rows = [ws.row_values(i) for i in range(ws.nrows)]
    else:
        wb = openpyxl.load_workbook(path, read_only=True)
        ws = wb.active
        all_rows = [list(r) for r in ws.iter_rows(values_only=True)]

    if not all_rows:
        raise RuntimeError('Excel为空')

    headers = [str(h).strip() if h is not None else '' for h in all_rows[0]]
    records = []
    for row in all_rows[1:]:
        if any(v is not None and str(v).strip() != '' for v in row):
            first = str(row[0]) if row[0] is not None else ''
            if '示例' in first or '↓' in first:
                continue
            records.append({headers[i]: (str(v) if v is not None else '') for i, v in enumerate(row)})
    return headers, records


def data_file_path(kind):
    filename = VEHICLE_DATA_FILE if kind == 'vehicle' else DRIVER_DATA_FILE
    return os.path.join(app.config['DATA_FOLDER'], filename)


def invalidate_data_cache(kind):
    if kind == 'vehicle':
        DATA_CACHE['vehicle_mtime'] = None
        DATA_CACHE['vehicle_records'] = []
    else:
        DATA_CACHE['driver_mtime'] = None
        DATA_CACHE['driver_records'] = []


def read_managed_workbook(kind):
    path = data_file_path(kind)
    if not os.path.exists(path):
        raise RuntimeError('资料文件不存在')
    wb = openpyxl.load_workbook(path)
    ws = wb.active
    headers = [compact_text(ws.cell(row=1, column=i).value) for i in range(1, ws.max_column + 1)]
    return path, wb, ws, headers


def row_values(ws, headers, row_num):
    return {
        header: compact_text(ws.cell(row=row_num, column=i + 1).value)
        for i, header in enumerate(headers)
        if header
    }


def managed_search(kind, query):
    query = compact_text(query)
    if len(query) < 1:
        return [], []

    path, wb, ws, headers = read_managed_workbook(kind)
    items = []
    for row_num in range(2, ws.max_row + 1):
        values = row_values(ws, headers, row_num)
        if not any(values.values()):
            continue
        first = next(iter(values.values()), '')
        if '示例' in first or '↓' in first:
            continue

        if kind == 'vehicle':
            plate = values.get('车牌号码', '')
            rank = rank_plate_match(query, plate, normalize_plate_query(plate))
            sort_key = normalize_plate_query(plate)
        else:
            rank = rank_match(query, values.get('姓名', ''))
            sort_key = normalize_query(values.get('姓名', ''))
        if rank is None:
            continue
        items.append({
            'id': row_num,
            'rank': rank,
            'sort_key': sort_key,
            'values': values,
        })

    items.sort(key=lambda item: (item['rank'], item['sort_key'], item['id']))
    for item in items:
        item.pop('rank', None)
        item.pop('sort_key', None)
    return headers, items[:50]


def backup_data_file(path):
    import shutil
    backup_dir = os.path.join(app.config['DATA_FOLDER'], 'backups')
    os.makedirs(backup_dir, exist_ok=True)
    backup_name = os.path.basename(path) + '.bak.' + uuid.uuid4().hex[:8]
    shutil.copy2(path, os.path.join(backup_dir, backup_name))


def update_managed_row(kind, row_id, values):
    try:
        row_num = int(row_id)
    except Exception:
        raise RuntimeError('无效记录ID')

    if not isinstance(values, dict):
        raise RuntimeError('提交数据格式错误')

    path, wb, ws, headers = read_managed_workbook(kind)
    if row_num < 2 or row_num > ws.max_row:
        raise RuntimeError('记录不存在')

    header_to_col = {header: i + 1 for i, header in enumerate(headers) if header}
    allowed_values = {k: compact_text(v) for k, v in values.items() if k in header_to_col}
    if not allowed_values:
        raise RuntimeError('没有可更新字段')

    backup_data_file(path)
    for header, value in allowed_values.items():
        ws.cell(row=row_num, column=header_to_col[header]).value = value
    wb.save(path)
    os.chmod(path, 0o600)
    invalidate_data_cache(kind)
    return row_values(ws, headers, row_num)


def compact_text(value):
    if value is None:
        return ''
    text = str(value).strip()
    if text.endswith('.0'):
        text = text[:-2]
    return text


def normalize_query(value):
    return re.sub(r'\s+', '', compact_text(value)).lower()


def normalize_plate_query(value):
    text = normalize_query(value).upper()
    for prefix in ('苏K-', '苏K', 'K-', 'K'):
        if text.startswith(prefix):
            text = text[len(prefix):]
            break
    if '(' in text:
        text = text.split('(', 1)[0]
    return text


def load_records(kind):
    if kind == 'vehicle':
        filename = VEHICLE_DATA_FILE
        mtime_key = 'vehicle_mtime'
        records_key = 'vehicle_records'
    else:
        filename = DRIVER_DATA_FILE
        mtime_key = 'driver_mtime'
        records_key = 'driver_records'

    path = os.path.join(app.config['DATA_FOLDER'], filename)
    if not os.path.exists(path):
        return []

    mtime = os.path.getmtime(path)
    if DATA_CACHE[mtime_key] == mtime:
        return DATA_CACHE[records_key]

    headers, rows = read_excel(path)
    records = []
    for row in rows:
        if kind == 'vehicle':
            plate = compact_text(row.get('车牌号码'))
            if not plate:
                continue
            records.append({
                'plate': plate,
                'plate_tail': normalize_plate_query(plate),
                'seat_count': compact_text(row.get('核载')),
                'car_type': compact_text(row.get('车辆类型')),
                'transport_cert': compact_text(row.get('道路运输证字号')),
            })
        else:
            name = compact_text(row.get('姓名'))
            if not name:
                continue
            records.append({
                'name': name,
                'phone': compact_text(row.get('手机')),
                'cert': compact_text(row.get('从业资格证号')),
            })

    DATA_CACHE[mtime_key] = mtime
    DATA_CACHE[records_key] = records
    logger.info('Loaded %s records: %d', kind, len(records))
    return records


def rank_match(query, value):
    q = normalize_query(query)
    v = normalize_query(value)
    if not q or not v:
        return None
    if v == q:
        return 0
    if v.startswith(q):
        return 1
    if q in v:
        return 2
    return None


def rank_plate_match(query, plate, plate_tail):
    q = normalize_plate_query(query)
    if not q:
        return None
    candidates = [normalize_plate_query(plate), normalize_plate_query(plate_tail)]
    if q in candidates:
        return 0
    if any(c.startswith(q) for c in candidates):
        return 1
    if any(q in c for c in candidates):
        return 2
    return None


def first_vehicle_match(query):
    matches = []
    for record in load_records('vehicle'):
        rank = rank_plate_match(query, record.get('plate', ''), record.get('plate_tail', ''))
        if rank is not None:
            matches.append((rank, record.get('plate', ''), record))
    if not matches:
        return None
    matches.sort(key=lambda item: (item[0], item[1]))
    return matches[0][2]


def first_driver_match(query):
    matches = []
    for record in load_records('driver'):
        rank = rank_match(query, record.get('name', ''))
        if rank is not None:
            matches.append((rank, record.get('name', ''), record))
    if not matches:
        return None
    matches.sort(key=lambda item: (item[0], item[1]))
    return matches[0][2]


def fill_if_blank(data, key, value):
    if not compact_text(data.get(key)) and compact_text(value):
        data[key] = compact_text(value)


def enrich_fill_data(data):
    plate = compact_text(data.get('plate'))
    if plate:
        vehicle = first_vehicle_match(plate)
        if vehicle:
            data['plate'] = vehicle.get('plate_tail') or normalize_plate(vehicle.get('plate', ''))
            fill_if_blank(data, 'seat_count', vehicle.get('seat_count'))
            fill_if_blank(data, 'car_type', vehicle.get('car_type'))
            fill_if_blank(data, 'transport_cert', vehicle.get('transport_cert'))

    for index in (1, 2):
        name_key = 'driver%d_name' % index
        name = compact_text(data.get(name_key))
        if not name:
            continue
        driver = first_driver_match(name)
        if driver:
            data[name_key] = driver.get('name') or name
            fill_if_blank(data, 'driver%d_phone' % index, driver.get('phone'))
            fill_if_blank(data, 'driver%d_cert' % index, driver.get('cert'))
    return data


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def allowed_template_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'doc', 'docx'}


def template_display_name(filename):
    name = os.path.splitext(filename)[0]
    return name.replace('-电子合同', '').replace('电子合同', '').strip('-_ ')


def list_builtin_templates():
    folder = app.config['TEMPLATE_FOLDER']
    if not os.path.isdir(folder):
        return []
    items = []
    for filename in sorted(os.listdir(folder)):
        if not allowed_template_file(filename):
            continue
        path = os.path.join(folder, filename)
        if not os.path.isfile(path):
            continue
        items.append({
            'id': filename,
            'name': template_display_name(filename),
            'filename': filename,
        })
    return items


def get_builtin_template_path(template_id):
    template_id = compact_text(template_id)
    if not template_id or '/' in template_id or '\\' in template_id or not allowed_template_file(template_id):
        raise RuntimeError('无效模板')
    path = os.path.join(app.config['TEMPLATE_FOLDER'], template_id)
    if not os.path.isfile(path):
        raise RuntimeError('模板不存在')
    return path


def fetch_json(url, params):
    full_url = url + '?' + urlencode(params)
    with urlopen(full_url, timeout=8) as resp:
        body = resp.read().decode('utf-8')
    return json.loads(body)


def amap_geocode(address):
    key = app.config['AMAP_API_KEY']
    if not key:
        raise RuntimeError('服务器未配置高德地图Key，请配置 AMAP_API_KEY')
    result = fetch_json('https://restapi.amap.com/v3/geocode/geo', {
        'key': key,
        'address': address,
        'output': 'json',
    })
    if result.get('status') != '1':
        raise RuntimeError(result.get('info') or '地址解析失败')
    geocodes = result.get('geocodes') or []
    if not geocodes:
        raise RuntimeError('未找到地址：' + address)
    item = geocodes[0]
    location = item.get('location')
    if not location:
        raise RuntimeError('地址缺少坐标：' + address)
    return {
        'location': location,
        'formatted_address': item.get('formatted_address') or address,
    }


def amap_route_between(origin_geo, dest_geo, waypoint_geos=None):
    waypoint_geos = waypoint_geos or []
    key = app.config['AMAP_API_KEY']
    params = {
        'key': key,
        'origin': origin_geo['location'],
        'destination': dest_geo['location'],
        'strategy': 0,
        'output': 'json',
    }
    if waypoint_geos:
        params['waypoints'] = ';'.join(item['location'] for item in waypoint_geos)
    result = fetch_json('https://restapi.amap.com/v3/direction/driving', params)
    if result.get('status') != '1':
        raise RuntimeError(result.get('info') or '路线规划失败')
    paths = ((result.get('route') or {}).get('paths') or [])
    if not paths:
        raise RuntimeError('未找到驾车路线')
    path = paths[0]
    distance_m = int(float(path.get('distance') or 0))
    duration_s = int(float(path.get('duration') or 0))
    return distance_m, duration_s


def amap_driving_distance(origin, destination, waypoints=None):
    waypoints = waypoints or []
    origin_geo = amap_geocode(origin)
    dest_geo = amap_geocode(destination)
    waypoint_geos = [amap_geocode(item) for item in waypoints if compact_text(item)]
    distance_m, duration_s = amap_route_between(origin_geo, dest_geo, waypoint_geos)

    nodes = [origin_geo] + waypoint_geos + [dest_geo]
    segments = []
    for idx in range(len(nodes) - 1):
        seg_distance_m, seg_duration_s = amap_route_between(nodes[idx], nodes[idx + 1])
        segments.append({
            'from': nodes[idx]['formatted_address'],
            'to': nodes[idx + 1]['formatted_address'],
            'distance_m': seg_distance_m,
            'distance_km': round(seg_distance_m / 1000, 1),
            'duration_min': round(seg_duration_s / 60),
        })

    return {
        'provider': 'amap',
        'origin': origin_geo['formatted_address'],
        'destination': dest_geo['formatted_address'],
        'waypoints': [item['formatted_address'] for item in waypoint_geos],
        'segments': segments,
        'distance_m': distance_m,
        'distance_km': round(distance_m / 1000, 1),
        'duration_min': round(duration_s / 60),
    }


def is_logged_in():
    return session.get('logged_in') is True


def session_expired():
    last_active = session.get('last_active')
    if not last_active:
        return False
    return time.time() - float(last_active) > IDLE_TIMEOUT_SECONDS


def client_key():
    return request.headers.get('X-Real-IP') or request.remote_addr or 'unknown'


def login_locked(key):
    now = time.time()
    attempts = [t for t in LOGIN_ATTEMPTS.get(key, []) if now - t < LOGIN_LOCK_SECONDS]
    LOGIN_ATTEMPTS[key] = attempts
    return len(attempts) >= LOGIN_MAX_ATTEMPTS


def record_login_failure(key):
    now = time.time()
    attempts = [t for t in LOGIN_ATTEMPTS.get(key, []) if now - t < LOGIN_LOCK_SECONDS]
    attempts.append(now)
    LOGIN_ATTEMPTS[key] = attempts


def clear_login_failures(key):
    LOGIN_ATTEMPTS.pop(key, None)


@app.before_request
def require_login():
    allowed_endpoints = {'login', 'static'}
    if request.endpoint in allowed_endpoints:
        return None
    if is_logged_in() and not session_expired():
        session['last_active'] = time.time()
        return None
    session.clear()
    if request.path.startswith('/api/'):
        return jsonify({'error': '登录已过期，请重新登录'}), 401
    return redirect(url_for('login', next=request.full_path if request.query_string else request.path))


@app.after_request
def add_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['Referrer-Policy'] = 'same-origin'
    if request.path.startswith('/api/lookup') or request.path.startswith('/api/manage'):
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        response.headers['Pragma'] = 'no-cache'
    return response


def normalize_plate(val):
    """车牌号去掉苏K-前缀"""
    if val and val.upper().startswith('苏K-'):
        return val[3:]
    return val


def fill_doc_template(template_path, data, output_path):
    import re as _re
    from docx import Document
    from collections import defaultdict

    def copy_run_style(src, dst):
        if not src:
            return
        dst.bold = src.bold
        dst.italic = src.italic
        dst.font.name = src.font.name
        dst.font.size = src.font.size
        dst.font.color.rgb = src.font.color.rgb

    def rebuild_para_with_segments(para, segments):
        template_run = para.runs[0] if para.runs else None
        for run in para.runs:
            run.text = ''
        first = True
        for text, underline in segments:
            if text == '':
                continue
            if first and para.runs:
                run = para.runs[0]
                first = False
            else:
                run = para.add_run()
                copy_run_style(template_run, run)
            run.text = text
            if underline:
                run.underline = True

    def placeholder_segments(text):
        pattern = _re.compile(r'\{\{([A-Za-z0-9_]+)\}\}')
        segments = []
        pos = 0
        for match in pattern.finditer(text):
            if match.start() > pos:
                segments.append((text[pos:match.start()], False))
            val = data.get(match.group(1), '')
            if val:
                segments.append((str(val), True))
            pos = match.end()
        if pos < len(text):
            segments.append((text[pos:], False))
        return segments

    doc = Document(template_path)
    all_parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_parts.extend(p.text for p in cell.paragraphs)
    all_text = ' '.join(all_parts)
    use_placeholder = '{{' in all_text

    if use_placeholder:
        def replace_in_para(para):
            full = ''.join(r.text for r in para.runs)
            if '{{' in full:
                rebuild_para_with_segments(para, placeholder_segments(full))
        for para in doc.paragraphs:
            replace_in_para(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_para(para)
    else:
        def replace_blanks_in_para(para, replacements):
            full = ''.join(r.text for r in para.runs)
            gaps = list(_re.finditer(r' {2,}', full))
            marks = []
            for gap_i, val in replacements.items():
                if gap_i < len(gaps):
                    m = gaps[gap_i]
                    marks.append((m.start(), m.end(), str(val)))
            if not marks:
                return
            marks.sort()
            segments = []
            pos = 0
            for start, end, val in marks:
                if start > pos:
                    segments.append((full[pos:start], False))
                segments.append((val, True))
                pos = end
            if pos < len(full):
                segments.append((full[pos:], False))
            rebuild_para_with_segments(para, segments)

        para_replacements = defaultdict(dict)
        for para_i, gap_i, key in NANJING_BLANK_MAP:
            val = data.get(key, '')
            if val:
                para_replacements[para_i][gap_i] = val

        for para_i, replacements in para_replacements.items():
            if para_i < len(doc.paragraphs):
                replace_blanks_in_para(doc.paragraphs[para_i], replacements)

    _fix_to_single_page(doc)
    doc.save(output_path)


def _fix_to_single_page(doc):
    """尽量保留原Word/WPS排版，不强制改字体、字号或页边距。"""
    return


def run_cmd(cmd, timeout=60):
    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    try:
        stdout, stderr = proc.communicate(timeout=timeout)
    except subprocess.TimeoutExpired:
        proc.kill()
        proc.communicate()
        raise RuntimeError('命令超时: ' + ' '.join(cmd))
    return proc.returncode, stdout.decode('utf-8', errors='replace'), stderr.decode('utf-8', errors='replace')


def _is_stamp_red(r, g, b):
    return r > 115 and r - g > 22 and r - b > 18 and g < 225 and b < 225


def _find_red_stamp_boxes(img):
    """找出页面底部的红色印章区域，用于避免最终压缩导致印章变形。"""
    pix = img.load()
    width, height = img.size
    mask = set()
    for y in range(height):
        for x in range(width):
            if _is_stamp_red(*pix[x, y]):
                mask.add((x, y))

    seen = set()
    comps = []
    for point in list(mask):
        if point in seen:
            continue
        stack = [point]
        seen.add(point)
        xs = []
        ys = []
        while stack:
            x, y = stack.pop()
            xs.append(x)
            ys.append(y)
            for nx, ny in ((x + 1, y), (x - 1, y), (x, y + 1), (x, y - 1)):
                if (nx, ny) in mask and (nx, ny) not in seen:
                    seen.add((nx, ny))
                    stack.append((nx, ny))
        if len(xs) > 150:
            comps.append((min(xs), min(ys), max(xs) + 1, max(ys) + 1, len(xs)))

    bottom = [c for c in comps if c[1] > height * 0.75]
    stamp_boxes = []
    for group in (
        [c for c in bottom if (c[0] + c[2]) / 2 < width / 2],
        [c for c in bottom if (c[0] + c[2]) / 2 >= width / 2],
    ):
        if not group:
            continue
        group = sorted(group, key=lambda c: c[4], reverse=True)[:20]
        stamp_boxes.append((
            min(c[0] for c in group),
            min(c[1] for c in group),
            max(c[2] for c in group),
            max(c[3] for c in group),
        ))
    return stamp_boxes


def _print_page_with_round_stamps(source):
    """把长内容压进打印页画布，同时保持红色印章按等比例显示。"""
    from PIL import Image

    page_w, page_h = PRINT_PAGE_SIZE
    box_left, box_top, box_right, box_bottom = PRINT_CONTENT_BOX
    box_w = box_right - box_left
    box_h = box_bottom - box_top

    canvas = Image.new('RGB', PRINT_PAGE_SIZE, (255, 255, 255))
    canvas.paste(source.resize((box_w, box_h), Image.LANCZOS), (box_left, box_top))

    src_w, src_h = source.size
    scale_x = box_w / src_w
    scale_y = box_h / src_h
    canvas_pix = canvas.load()
    erase_top = max(0, box_top + int(src_h * 0.75 * scale_y) - 80)
    erase_bottom = min(page_h, box_bottom + 120)
    for y in range(erase_top, erase_bottom):
        for x in range(box_left, box_right):
            if _is_stamp_red(*canvas_pix[x, y]):
                canvas_pix[x, y] = (255, 255, 255)

    for x1, y1, x2, y2 in _find_red_stamp_boxes(source):
        pad = 18
        x1 = max(0, x1 - pad)
        y1 = max(0, y1 - pad)
        x2 = min(src_w, x2 + pad)
        y2 = min(src_h, y2 + pad)
        crop = source.crop((x1, y1, x2, y2)).convert('RGBA')
        data = []
        for r, g, b, a in crop.getdata():
            if _is_stamp_red(r, g, b):
                data.append((r, g, b, 255))
            else:
                data.append((255, 255, 255, 0))
        crop.putdata(data)

        # 使用横向缩放比例回贴红章层，避免整页纵向压缩把圆章压成椭圆。
        stamp_w = max(1, round(crop.width * scale_x))
        stamp_h = max(1, round(crop.height * scale_x))
        crop = crop.resize((stamp_w, stamp_h), Image.LANCZOS)
        paste_x = box_left + round(x1 * scale_x)
        paste_y = box_top + round(y1 * scale_y) - round((stamp_h - (y2 - y1) * scale_y) / 2)
        canvas.paste(crop, (paste_x, paste_y), crop)

    return canvas


def docx_to_single_image(docx_path, output_dir, out_name='contract'):
    """docx -> odt(修改行距) -> PDF -> 按Word打印页比例输出PNG"""
    import zipfile, re

    # 先把 docx 转 odt
    code_odt, _, err_odt = run_cmd(
        ['libreoffice', '--headless', '--convert-to', 'odt', '--outdir', output_dir, docx_path]
    )
    odt_name = os.path.splitext(os.path.basename(docx_path))[0] + '.odt'
    odt_path = os.path.join(output_dir, odt_name)

    if code_odt == 0 and os.path.exists(odt_path):
        # 在 odt 里直接压缩段落行距，贴近Word/WPS打印一页的视觉密度
        fixed_odt = odt_path.replace('.odt', '_fixed.odt')
        try:
            with zipfile.ZipFile(odt_path) as zin:
                content_xml = zin.read('content.xml').decode()
                styles_xml = zin.read('styles.xml').decode()

            def inject_line_height(xml):
                def replacer(m):
                    tag = m.group(0)
                    if 'fo:line-height' not in tag:
                        tag = tag.replace('<style:paragraph-properties',
                                         '<style:paragraph-properties fo:line-height="55%"')
                    return tag
                xml = re.sub(r'<style:paragraph-properties[^>]*>', replacer, xml)
                xml = re.sub(r'fo:margin-top="[^"]*"', 'fo:margin-top="0cm"', xml)
                xml = re.sub(r'fo:margin-bottom="[^"]*"', 'fo:margin-bottom="0cm"', xml)
                return xml

            content_xml = inject_line_height(content_xml)
            styles_xml = inject_line_height(styles_xml)

            with zipfile.ZipFile(fixed_odt, 'w', zipfile.ZIP_DEFLATED) as zout:
                with zipfile.ZipFile(odt_path) as zin:
                    for item in zin.namelist():
                        if item == 'content.xml':
                            zout.writestr(item, content_xml.encode())
                        elif item == 'styles.xml':
                            zout.writestr(item, styles_xml.encode())
                        else:
                            zout.writestr(item, zin.read(item))

            src_for_pdf = fixed_odt
        except Exception:
            src_for_pdf = docx_path  # 降级回原始docx
    else:
        src_for_pdf = docx_path

    code, out, err = run_cmd(
        ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, src_for_pdf]
    )
    if code != 0:
        raise RuntimeError('LibreOffice转换失败: ' + err)

    pdf_name = os.path.splitext(os.path.basename(src_for_pdf))[0] + '.pdf'
    pdf_path = os.path.join(output_dir, pdf_name)
    if not os.path.exists(pdf_path):
        raise RuntimeError('PDF文件未生成')

    img_prefix = os.path.join(output_dir, 'page')
    code2, out2, err2 = run_cmd(
        ['pdftoppm', '-r', '150', '-png', pdf_path, img_prefix]
    )
    if code2 != 0:
        raise RuntimeError('pdftoppm转换失败: ' + err2)

    pages = sorted([
        os.path.join(output_dir, f)
        for f in os.listdir(output_dir)
        if f.startswith('page') and f.endswith('.png')
    ])
    if not pages:
        raise RuntimeError('未找到生成的图片')

    from PIL import Image, ImageChops

    MARGIN = 30  # 约5mm @150dpi

    def get_content_bounds(img):
        diff = ImageChops.difference(img, Image.new('RGB', img.size, (255, 255, 255)))
        bb = diff.getbbox()
        return bb

    imgs = [Image.open(p).convert('RGB') for p in pages]

    if len(imgs) == 1:
        bb = get_content_bounds(imgs[0])
        bottom = min(imgs[0].height, bb[3] + MARGIN) if bb else imgs[0].height
        final = imgs[0].crop((0, 0, imgs[0].width, bottom))
    else:
        # 第1页：从0到内容底部+边距（裁底部空白）
        # 中间页/最后页：裁顶部分页空白（去掉page break空白），底部+边距
        crops = []
        for idx, img in enumerate(imgs):
            bb = get_content_bounds(img)
            if not bb:
                continue
            is_first = (idx == 0)
            top = 0 if is_first else max(0, bb[1] - MARGIN)
            bottom = min(img.height, bb[3] + MARGIN)
            crops.append(img.crop((0, top, img.width, bottom)))

        if not crops:
            raise RuntimeError('转换后的图片为空')

        w = max(c.width for c in crops)
        h = sum(c.height for c in crops)
        merged = Image.new('RGB', (w, h), (255, 255, 255))
        y = 0
        for c in crops:
            merged.paste(c, ((w - c.width) // 2, y))
            y += c.height
        final = merged

    if final.height / final.width > 1.6:
        final = _print_page_with_round_stamps(final)

    final_path = os.path.join(output_dir, out_name + '.png')
    final.save(final_path)
    return final_path


def split_date(data, date_key, year_key, month_key, day_key):
    """把 yyyy-mm-dd 日期字段拆分为年月日三个字段"""
    val = data.pop(date_key, '') or ''
    val = val.strip()
    if val and '-' in val:
        parts = val.split('-')
        if len(parts) == 3:
            data[year_key]  = parts[0].lstrip('0') or parts[0]
            data[month_key] = parts[1].lstrip('0') or parts[1]
            data[day_key]   = parts[2].lstrip('0') or parts[2]
            return
    if not val and (data.get(year_key) or data.get(month_key) or data.get(day_key)):
        return
    # 没填或格式不对，留空
    data[year_key] = data[month_key] = data[day_key] = ''


def build_fill_data(source_data):
    """处理填充数据：按车牌/司机姓名补全资料、拆分日期、归一化车牌"""
    data = dict(source_data)
    data = enrich_fill_data(data)
    data['plate'] = normalize_plate(data.get('plate', ''))
    split_date(data, 'start_date', 'start_year', 'start_month', 'start_day')
    split_date(data, 'end_date',   'end_year',   'end_month',   'end_day')
    split_date(data, 'sign_date',  'sign_year',  'sign_month',  'sign_day')
    return data


def prepare_template_file(tmpl_file, task_id, template_id=''):
    if tmpl_file:
        tmpl_ext = tmpl_file.filename.rsplit('.', 1)[1].lower() if '.' in tmpl_file.filename else 'docx'
        tmpl_path = os.path.join(app.config['UPLOAD_FOLDER'], task_id + '_template.' + tmpl_ext)
        tmpl_file.save(tmpl_path)
    else:
        source_path = get_builtin_template_path(template_id)
        tmpl_ext = source_path.rsplit('.', 1)[1].lower()
        tmpl_path = os.path.join(app.config['UPLOAD_FOLDER'], task_id + '_template.' + tmpl_ext)
        shutil.copy2(source_path, tmpl_path)


    if tmpl_ext == 'doc':
        code, out, err = run_cmd(
            ['libreoffice', '--headless', '--convert-to', 'docx',
             '--outdir', app.config['UPLOAD_FOLDER'], tmpl_path]
        )
        tmpl_docx = tmpl_path.replace('.doc', '.docx')
        if not os.path.exists(tmpl_docx):
            raise RuntimeError('doc转docx失败: ' + err)
        tmpl_path = tmpl_docx
    return tmpl_path


def generate_contract_images(tmpl_path, rows, task_id):
    task_dir = os.path.join(app.config['OUTPUT_FOLDER'], task_id)
    os.makedirs(task_dir)

    if len(rows) == 1:
        row_dir = os.path.join(task_dir, 'row_1')
        os.makedirs(row_dir)
        filled_docx = os.path.join(row_dir, 'contract.docx')
        fill_doc_template(tmpl_path, build_fill_data(rows[0]), filled_docx)
        return {
            'single_image': docx_to_single_image(filled_docx, row_dir, 'contract')
        }

    zip_path = os.path.join(app.config['OUTPUT_FOLDER'], task_id + '.zip')
    errors = []
    generated = 0

    with zipfile.ZipFile(zip_path, 'w') as zf:
        for idx, row in enumerate(rows):
            row_dir = os.path.join(task_dir, 'row_%d' % (idx + 1))
            os.makedirs(row_dir)
            try:
                filled_docx = os.path.join(row_dir, 'contract.docx')
                fill_doc_template(tmpl_path, build_fill_data(row), filled_docx)
                img_path = docx_to_single_image(filled_docx, row_dir, 'contract')
                zf.write(img_path, '合同_%d.png' % (idx + 1))
                generated += 1
            except Exception as e:
                errors.append('第%d行: %s' % (idx + 1, str(e)))
                logger.error('Row %d failed: %s', idx + 1, e)

    if generated == 0:
        raise RuntimeError(json.dumps({'error': '所有记录生成失败', 'details': errors}, ensure_ascii=False))

    return {
        'task_id': task_id,
        'generated': generated,
        'errors': errors,
        'download_url': '/api/download/' + task_id
    }


@app.route('/')
def index():
    return render_template('index.html', fields=FIELD_DEFS)


@app.route('/login', methods=['GET', 'POST'])
def login():
    error = ''
    if request.method == 'POST':
        key = client_key()
        if login_locked(key):
            return render_template('login.html', error='登录失败次数过多，请稍后再试', next_url=request.args.get('next', '')), 429
        username = request.form.get('username', '')
        password = request.form.get('password', '')
        if hmac.compare_digest(username, LOGIN_USERNAME) and hmac.compare_digest(password, LOGIN_PASSWORD):
            clear_login_failures(key)
            session.clear()
            session.permanent = True
            session['logged_in'] = True
            session['username'] = LOGIN_USERNAME
            session['last_active'] = time.time()
            next_url = request.form.get('next') or url_for('index')
            if not next_url.startswith('/'):
                next_url = url_for('index')
            return redirect(next_url)
        record_login_failure(key)
        error = '账号或密码错误'
    return render_template('login.html', error=error, next_url=request.args.get('next', ''))


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


@app.route('/api/fields')
def get_fields():
    return jsonify(FIELD_DEFS)


@app.route('/api/templates')
def get_templates():
    return jsonify({'items': list_builtin_templates()})


@app.route('/api/distance', methods=['POST'])
def calculate_distance():
    payload = request.get_json(silent=True) or {}
    origin = compact_text(payload.get('origin'))
    destination = compact_text(payload.get('destination'))
    waypoints = payload.get('waypoints') or []
    if not isinstance(waypoints, list):
        return jsonify({'error': '途经地格式错误'}), 400
    waypoints = [compact_text(item) for item in waypoints if compact_text(item)]
    if not origin or not destination:
        return jsonify({'error': '请输入出发地和目的地'}), 400
    try:
        return jsonify(amap_driving_distance(origin, destination, waypoints))
    except Exception as e:
        logger.error('calculate_distance failed: %s', e)
        return jsonify({'error': str(e)}), 400


@app.route('/api/lookup/vehicle')
def lookup_vehicle():
    query = request.args.get('q', '').strip()
    if len(query) < 1:
        return jsonify({'items': []})

    matches = []
    for record in load_records('vehicle'):
        rank = rank_plate_match(query, record.get('plate', ''), record.get('plate_tail', ''))
        if rank is None:
            continue
        matches.append((rank, record.get('plate', ''), record))

    matches.sort(key=lambda item: (item[0], item[1]))
    items = []
    for _, _, record in matches[:8]:
        items.append({
            'plate': record.get('plate', ''),
            'plate_tail': record.get('plate_tail', ''),
            'seat_count': record.get('seat_count', ''),
            'car_type': record.get('car_type', ''),
            'transport_cert': record.get('transport_cert', ''),
        })
    return jsonify({'items': items})


@app.route('/api/lookup/driver')
def lookup_driver():
    query = request.args.get('q', '').strip()
    if len(query) < 1:
        return jsonify({'items': []})

    matches = []
    for record in load_records('driver'):
        rank = rank_match(query, record.get('name', ''))
        if rank is None:
            continue
        matches.append((rank, record.get('name', ''), record))

    matches.sort(key=lambda item: (item[0], item[1]))
    items = []
    for _, _, record in matches[:8]:
        items.append({
            'name': record.get('name', ''),
            'phone': record.get('phone', ''),
            'cert': record.get('cert', ''),
        })
    return jsonify({'items': items})


@app.route('/api/manage/drivers')
def manage_drivers():
    query = request.args.get('q', '').strip()
    try:
        headers, items = managed_search('driver', query)
        return jsonify({'headers': headers, 'items': items})
    except Exception as e:
        logger.error('manage_drivers failed: %s', e)
        return jsonify({'error': str(e)}), 500


@app.route('/api/manage/drivers/<row_id>', methods=['POST'])
def update_driver(row_id):
    payload = request.get_json(silent=True) or {}
    try:
        values = update_managed_row('driver', row_id, payload.get('values', {}))
        return jsonify({'success': True, 'values': values})
    except Exception as e:
        logger.error('update_driver failed: %s', e)
        return jsonify({'error': str(e)}), 400


@app.route('/api/manage/vehicles')
def manage_vehicles():
    query = request.args.get('q', '').strip()
    try:
        headers, items = managed_search('vehicle', query)
        return jsonify({'headers': headers, 'items': items})
    except Exception as e:
        logger.error('manage_vehicles failed: %s', e)
        return jsonify({'error': str(e)}), 500


@app.route('/api/manage/vehicles/<row_id>', methods=['POST'])
def update_vehicle(row_id):
    payload = request.get_json(silent=True) or {}
    try:
        values = update_managed_row('vehicle', row_id, payload.get('values', {}))
        return jsonify({'success': True, 'values': values})
    except Exception as e:
        logger.error('update_vehicle failed: %s', e)
        return jsonify({'error': str(e)}), 400


@app.route('/api/parse-excel', methods=['POST'])
def parse_excel():
    if 'excel' not in request.files:
        return jsonify({'error': '请上传Excel文件'}), 400
    f = request.files['excel']
    if not allowed_file(f.filename):
        return jsonify({'error': '不支持的文件格式'}), 400
    ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
    path = os.path.join(app.config['UPLOAD_FOLDER'], str(uuid.uuid4())[:8] + '.' + ext)
    f.save(path)
    try:
        headers, data = read_excel(path)
        return jsonify({'headers': headers, 'preview': data[:5], 'total': len(data)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-manual', methods=['POST'])
def generate_manual():
    """手动输入单条数据，生成一张合同图片并直接返回"""
    tmpl_file = request.files.get('template')
    template_id = request.form.get('template_id', '')
    if not tmpl_file and not template_id:
        return jsonify({'error': '请选择或上传Word模板'}), 400

    raw_data = request.form.get('data')
    if not raw_data:
        return jsonify({'error': '请提供字段数据'}), 400

    try:
        fill_data = json.loads(raw_data)
    except Exception as e:
        return jsonify({'error': '数据格式错误: ' + str(e)}), 400

    task_id = str(uuid.uuid4())[:8]
    try:
        tmpl_path = prepare_template_file(tmpl_file, task_id, template_id)
        result = generate_contract_images(tmpl_path, [fill_data], task_id)
        return send_file(result['single_image'], mimetype='image/png',
                         as_attachment=True, download_name='合同.png')
    except Exception as e:
        logger.error('generate_manual failed: %s', e)
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-rows', methods=['POST'])
def generate_rows():
    """前端手动批量行生成，支持只填车牌和司机姓名后自动补全。"""
    tmpl_file = request.files.get('template')
    template_id = request.form.get('template_id', '')
    if not tmpl_file and not template_id:
        return jsonify({'error': '请选择或上传Word模板'}), 400

    raw_rows = request.form.get('rows')
    if not raw_rows:
        return jsonify({'error': '请提供批量数据'}), 400

    try:
        rows = json.loads(raw_rows)
    except Exception as e:
        return jsonify({'error': '批量数据格式错误: ' + str(e)}), 400

    if not isinstance(rows, list):
        return jsonify({'error': '批量数据必须是数组'}), 400

    clean_rows = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        clean = {k: compact_text(v) for k, v in row.items() if compact_text(v)}
        if clean:
            clean_rows.append(clean)

    if not clean_rows:
        return jsonify({'error': '请至少填写一行数据'}), 400

    task_id = str(uuid.uuid4())[:8]
    try:
        tmpl_path = prepare_template_file(tmpl_file, task_id, template_id)
        result = generate_contract_images(tmpl_path, clean_rows, task_id)
        if result.get('single_image'):
            return send_file(result['single_image'], mimetype='image/png',
                             as_attachment=True, download_name='合同.png')
        return jsonify(result)
    except RuntimeError as e:
        msg = str(e)
        try:
            parsed = json.loads(msg)
            return jsonify(parsed), 500
        except Exception:
            return jsonify({'error': msg}), 500
    except Exception as e:
        logger.error('generate_rows failed: %s', e)
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate', methods=['POST'])
def generate():
    """批量Excel生成，多份合同打包ZIP（每份一张图片）"""
    if 'excel' not in request.files:
        return jsonify({'error': '请上传Excel'}), 400

    tmpl_file = request.files.get('template')
    template_id = request.form.get('template_id', '')
    if not tmpl_file and not template_id:
        return jsonify({'error': '请选择或上传Word模板'}), 400
    excel_file = request.files['excel']
    mapping_raw = request.form.get('mapping')
    if not mapping_raw:
        return jsonify({'error': '请提供字段映射'}), 400
    mapping = json.loads(mapping_raw)

    task_id = str(uuid.uuid4())[:8]
    excel_ext = excel_file.filename.rsplit('.', 1)[-1].lower() if '.' in excel_file.filename else 'xlsx'
    excel_path = os.path.join(app.config['UPLOAD_FOLDER'], task_id + '_data.' + excel_ext)
    excel_file.save(excel_path)

    try:
        tmpl_path = prepare_template_file(tmpl_file, task_id, template_id)
        headers, records = read_excel(excel_path)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

    if not records:
        return jsonify({'error': 'Excel无数据行'}), 400

    rows = [{field_key: record.get(col_name, '') for field_key, col_name in mapping.items()} for record in records]
    try:
        result = generate_contract_images(tmpl_path, rows, task_id)
        if result.get('single_image'):
            return send_file(result['single_image'], mimetype='image/png',
                             as_attachment=True, download_name='合同.png')
        return jsonify(result)
    except RuntimeError as e:
        msg = str(e)
        try:
            parsed = json.loads(msg)
            return jsonify(parsed), 500
        except Exception:
            return jsonify({'error': msg}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/download/<task_id>')
def download(task_id):
    if not re.match(r'^[a-f0-9]{8}$', task_id):
        return jsonify({'error': '无效task_id'}), 400
    zip_path = os.path.join(app.config['OUTPUT_FOLDER'], task_id + '.zip')
    if not os.path.exists(zip_path):
        return jsonify({'error': '文件不存在'}), 404
    return send_file(zip_path, as_attachment=True, download_name='合同图片.zip')


if __name__ == '__main__':
    os.makedirs('uploads', exist_ok=True)
    os.makedirs('output', exist_ok=True)
    app.run(host='127.0.0.1', port=5000, debug=False)
